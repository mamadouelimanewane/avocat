from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_text(text):
    text = text.replace('**', '').replace('*', '')
    text = text.replace('’', "'").replace('“', '"').replace('”', '"').replace('…', '...')
    return re.sub(r'[^\x00-\x7F\xC0-\xFF]', '', text)

def create_docx(md_file, docx_file):
    doc = Document()
    
    # Stylistic setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Title
    title = doc.add_heading('LEXPREMIUM ERP - MANUEL COMPLET', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Chapters
    chapters = re.split(r'\n## ', content)
    
    for chapter in chapters[1:]:
        lines = chapter.strip().split('\n')
        if not lines: continue
        
        # Chapter Heading
        h = doc.add_heading(clean_text(lines[0]), level=1)
        h.runs[0].font.color.rgb = RGBColor(15, 23, 42) # Navy
        
        for line in lines[1:]:
            raw_line = line.strip()
            if not raw_line or raw_line == '---':
                continue
            
            cleaned = clean_text(raw_line)
            
            if raw_line.startswith('### '):
                p = doc.add_heading('Contenu', level=2)
            elif raw_line.startswith('• ') or raw_line.startswith('- '):
                p = doc.add_paragraph(cleaned[2:], style='List Bullet')
                # Optional: Bold the label before ':'
                if ':' in cleaned:
                    label, desc = cleaned[2:].split(':', 1)
                    p.clear()
                    run = p.add_run(label.strip() + ' : ')
                    run.bold = True
                    run.font.color.rgb = RGBColor(101, 67, 33) # Brown
                    p.add_run(desc.strip())
            else:
                p = doc.add_paragraph(cleaned)
                p.italic = True if "Le centre" in cleaned or " module " in cleaned else False

        doc.add_page_break()

    doc.save(docx_file)
    print(f"DOCX Generated: {docx_file}")

if __name__ == "__main__":
    create_docx("c:/gravity/avocat/MANUEL_UTILISATION_DETAILLE.md", "c:/gravity/avocat/release1.docx")
