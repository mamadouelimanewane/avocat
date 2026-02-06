from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx():
    doc = Document()
    
    # Title Page
    section = doc.sections[0]
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPOSITION COMMERCIALE & TECHNIQUE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(15, 23, 42) # Navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("LEXPREMIUM LITE - ÉDITION 2026")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(180, 140, 45) # Gold

    doc.add_paragraph("\n" * 5)
    
    # Vision
    doc.add_heading('1. VISION : LE CABINET NUMÉRIQUE DUO', level=1)
    p = doc.add_paragraph(
        "LexPremium Lite est l'édition haute performance conçue pour l'avocat moderne. "
        "Que vous soyez seul ou accompagné d'un assistant, le logiciel devient votre pilote automatique."
    )
    
    # Key Features
    doc.add_heading('2. FONCTIONNALITÉS CLÉS', level=1)
    
    features = [
        ("Tableau de Bord Dynamique", "Centre de commandement avec alertes forclusions et KPIs en temps réel."),
        ("Rédaction Assistée (LexAI)", "Génération d'actes (Assignations, Conclusions) aux normes OHADA."),
        ("Base de Connaissance IA", "Interrogez l'IA sur la jurisprudence et les textes de loi sénégalais."),
        ("Gestion Dossiers & Clients", "Archivage numérique complet et fiches clients à 360°."),
        ("Facturation & Débours", "Génération de mémoires d'honoraires et relances automatisées.")
    ]
    
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title} : ")
        run.bold = True
        p.add_run(desc)

    # Technical Specs
    doc.add_heading('3. SPÉCIFICATIONS TECHNIQUES', level=1)
    p = doc.add_paragraph("Le logiciel repose sur une infrastructure Cloud sécurisée :")
    doc.add_paragraph("Hébergement certifié conforme aux normes de protection des données.", style='List Bullet')
    doc.add_paragraph("Chiffrement AES-256 pour tous les documents et échanges.", style='List Bullet')
    doc.add_paragraph("Accessibilité multi-plateforme (iPhone, Android, PC/Mac).", style='List Bullet')

    # Conclusion
    doc.add_paragraph("\n")
    p = doc.add_paragraph("LexPremium : L'excellence technologique au service de votre expertise juridique.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(180, 140, 45)

    doc.save("c:/gravity/Avocat/avocat-lite/LexPremium_Lite_Proposition_2026.docx")
    print("DOCX Generated successfully.")

if __name__ == "__main__":
    generate_docx()
